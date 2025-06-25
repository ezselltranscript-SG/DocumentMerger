import os
import re
import shutil
import tempfile
import zipfile
import rarfile
import patoolib
from typing import List, Optional, Tuple
from fastapi import FastAPI, File, UploadFile, HTTPException, Form, Request
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
from pypdf import PdfWriter, PdfReader
import docx

app = FastAPI(title="PDF and DOCX Merger API")

# Configurar CORS para permitir solicitudes desde cualquier origen
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create directories if they don't exist
os.makedirs("uploads", exist_ok=True)

# Add CORS middleware to allow cross-origin requests
from fastapi.middleware.cors import CORSMiddleware

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)

# Helper function to extract part number from filename
def extract_part_number(filename: str) -> int:
    """Extract part number from filename (e.g., 'file_part1.pdf' -> 1)"""
    match = re.search(r'part(\d+)', filename.lower())
    if match:
        return int(match.group(1))
    return 0  # Default to 0 if no part number found

# Helper function to sort files by part number
def sort_files_by_part(files: List[UploadFile]) -> List[UploadFile]:
    """Sort files by their part number in the filename"""
    return sorted(files, key=lambda file: extract_part_number(file.filename))

# Helper function to merge PDF files
def merge_pdf_files(file_paths: List[str], output_path: str) -> None:
    """Merge multiple PDF files into a single PDF"""
    merger = PdfWriter()
    
    for file_path in file_paths:
        reader = PdfReader(file_path)
        for page in reader.pages:
            merger.add_page(page)
    
    with open(output_path, "wb") as output_file:
        merger.write(output_file)

# Helper function to extract files from compressed archives
def extract_compressed_file(file_path: str, extract_dir: str) -> List[str]:
    """Extract files from ZIP or RAR archive and return paths to extracted files"""
    extracted_files = []
    try:
        # Primero intentamos con zipfile
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)
                extracted_files = [os.path.join(extract_dir, name) for name in zip_ref.namelist() if not name.endswith('/')]
                return extracted_files
        except zipfile.BadZipFile:
            # No es un archivo ZIP, intentamos con RAR
            pass
        
        # Intentamos con rarfile
        try:
            rarfile.UNRAR_TOOL = 'unrar'
            with rarfile.RarFile(file_path) as rar_ref:
                rar_ref.extractall(extract_dir)
                extracted_files = [os.path.join(extract_dir, name) for name in rar_ref.namelist() if not rar_ref.getinfo(name).isdir()]
                return extracted_files
        except rarfile.NotRarFile:
            # No es un archivo RAR, intentamos con patoolib como último recurso
            pass
        
        # Último recurso: intentamos con patoolib que maneja varios formatos
        patoolib.extract_archive(file_path, outdir=extract_dir)
        for root, _, files in os.walk(extract_dir):
            for file in files:
                extracted_files.append(os.path.join(root, file))
        
        if not extracted_files:
            raise ValueError("No se pudieron extraer archivos del archivo comprimido")
            
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error extracting archive: {str(e)}")
    
    return extracted_files

# Helper function to filter files by extension
def filter_files_by_extension(file_paths: List[str], extensions: List[str]) -> List[str]:
    """Filter files by their extensions"""
    return [path for path in file_paths if os.path.splitext(path)[1].lower() in extensions]

# Helper function to merge DOCX files
def merge_docx_files_custom(file_paths: List[str], output_path: str) -> None:
    """Merge multiple DOCX files into a single DOCX preserving original format and headers"""
    import logging
    import shutil
    import tempfile
    import zipfile
    import os
    from docx import Document
    
    logging.info(f"Fusionando {len(file_paths)} archivos DOCX preservando formato original")
    logging.info(f"Archivos a fusionar: {[os.path.basename(f) for f in file_paths]}")

    if not file_paths:
        return

    if len(file_paths) == 1:
        shutil.copy(file_paths[0], output_path)
        return

    try:
        # Crear un directorio temporal para trabajar con los archivos
        with tempfile.TemporaryDirectory() as temp_dir:
            # Extraer el primer documento como base
            base_doc_path = file_paths[0]
            base_extract_dir = os.path.join(temp_dir, "base")
            os.makedirs(base_extract_dir, exist_ok=True)
            
            with zipfile.ZipFile(base_doc_path, 'r') as zip_ref:
                zip_ref.extractall(base_extract_dir)
            
            # Crear un nuevo documento combinado
            combined_dir = os.path.join(temp_dir, "combined")
            shutil.copytree(base_extract_dir, combined_dir)
            
            # Obtener el contenido del documento base
            with open(os.path.join(combined_dir, "word/document.xml"), 'r', encoding='utf-8') as f:
                base_content = f.read()
                
            # Encontrar la posición donde insertar el contenido de los otros documentos
            # (justo antes del cierre del cuerpo del documento)
            insert_pos = base_content.rfind("</w:body>")
            
            if insert_pos == -1:
                raise ValueError("No se pudo encontrar el final del cuerpo del documento base")
            
            # Preparar el contenido combinado
            combined_content = base_content[:insert_pos]
            
            # Agregar cada documento adicional
            for i, file_path in enumerate(file_paths[1:], 1):
                logging.info(f"Agregando documento {i}: {os.path.basename(file_path)}")
                
                # Extraer el documento actual
                doc_extract_dir = os.path.join(temp_dir, f"doc_{i}")
                os.makedirs(doc_extract_dir, exist_ok=True)
                
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    zip_ref.extractall(doc_extract_dir)
                
                # Leer el contenido del documento
                with open(os.path.join(doc_extract_dir, "word/document.xml"), 'r', encoding='utf-8') as f:
                    doc_content = f.read()
                
                # Extraer solo el contenido del cuerpo (entre <w:body> y </w:body>)
                body_start = doc_content.find("<w:body>") + len("<w:body>")
                body_end = doc_content.rfind("</w:body>")
                
                if body_start == -1 or body_end == -1:
                    raise ValueError(f"No se pudo extraer el cuerpo del documento {i}")
                
                body_content = doc_content[body_start:body_end]
                
                # Agregar un salto de página antes del contenido
                page_break = '<w:p><w:r><w:br w:type="page"/></w:r></w:p>'
                
                # Agregar el contenido al documento combinado
                combined_content += page_break + body_content
                
                # Copiar los estilos y relaciones del documento actual
                # (esto asegura que los encabezados y formatos se preserven)
                for item in os.listdir(os.path.join(doc_extract_dir, "word")):
                    if item.startswith("header") or item.startswith("footer") or item == "styles.xml":
                        src = os.path.join(doc_extract_dir, "word", item)
                        dst = os.path.join(combined_dir, "word", f"{i}_{item}")
                        
                        # Copiar el contenido del encabezado sin modificar
                        if item.startswith("header"):
                            with open(src, 'r', encoding='utf-8') as f:
                                header_content = f.read()
                            
                            with open(dst, 'w', encoding='utf-8') as f:
                                f.write(header_content)
                        else:
                            # Para otros archivos, simplemente copiar
                            shutil.copy(src, dst)
                        
                        # Actualizar el archivo de relaciones para incluir estos archivos
                        rels_file = os.path.join(combined_dir, "word", "_rels", "document.xml.rels")
                        if os.path.exists(rels_file):
                            with open(rels_file, 'r', encoding='utf-8') as f:
                                rels_content = f.read()
                            
                            # Agregar una nueva relación para este archivo
                            rel_id = f"rId{1000 + i}_{item.split('.')[0]}"
                            new_rel = f'<Relationship Id="{rel_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/header" Target="word/{i}_{item}"/>'
                            
                            # Añadir la nueva relación al contenido de relaciones
                            rels_content = rels_content.replace('</Relationships>', f'    {new_rel}\n</Relationships>')
                            
                            # Guardar el archivo de relaciones modificado
                            with open(rels_file, 'w', encoding='utf-8') as f:
                                f.write(rels_content)
            
            # Completar el documento combinado
            combined_content += "</w:body></w:document>"
            
            # Escribir el contenido combinado al archivo
            with open(os.path.join(combined_dir, "word/document.xml"), 'w', encoding='utf-8') as f:
                f.write(combined_content)
            
            # Crear un nuevo archivo ZIP (DOCX) con el contenido combinado
            with zipfile.ZipFile(output_path, 'w') as zip_out:
                for root, _, files in os.walk(combined_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arc_name = os.path.relpath(file_path, combined_dir)
                        zip_out.write(file_path, arc_name)
            
            logging.info(f"Documento combinado guardado en {output_path}")
    except Exception as e:
        logging.error(f"Error al fusionar documentos: {str(e)}")
        # Intentar con el enfoque alternativo si el principal falla
        try:
            logging.info("Intentando método alternativo de fusión...")
            result = merge_docx_files_simple(file_paths, output_path)
            return result
        except Exception as e2:
            logging.error(f"Error en método alternativo: {str(e2)}")
            raise HTTPException(status_code=500, detail=f"Error al fusionar documentos: {str(e)}")

def merge_docx_files_simple(file_paths: List[str], output_path: str) -> None:
    """Método alternativo más simple para fusionar documentos DOCX"""
    import logging
    import shutil
    import docx
    from docx import Document
    from docx.enum.text import WD_BREAK
    
    if not file_paths:
        return
        
    if len(file_paths) == 1:
        shutil.copy(file_paths[0], output_path)
        return
        
    try:
        # Crear un nuevo documento
        master = Document()
        
        # Para cada documento
        for i, path in enumerate(file_paths):
            doc = Document(path)
            
            # Si no es el primer documento, agregar un salto de página
            if i > 0:
                paragraph = master.add_paragraph()
                run = paragraph.add_run()
                run.add_break(WD_BREAK.PAGE)
            
            # Copiar cada párrafo
            for para in doc.paragraphs:
                p = master.add_paragraph()
                for run in para.runs:
                    r = p.add_run(run.text)
                    r.bold = run.bold
                    r.italic = run.italic
                    r.underline = run.underline
                    r.font.name = run.font.name
                    if run.font.size:
                        r.font.size = run.font.size
            
            # Copiar cada tabla
            for table in doc.tables:
                new_table = master.add_table(rows=len(table.rows), cols=len(table.columns))
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        if cell.text:
                            new_table.cell(i, j).text = cell.text
        
        # Guardar el documento combinado
        master.save(output_path)
        logging.info(f"Documento combinado guardado en {output_path} (método simple)")
        return True
    except Exception as e:
        logging.error(f"Error en método simple: {str(e)}")
        raise

@app.post("/api/merge/")
async def api_merge_files(
    file: Optional[UploadFile] = File(None),
    data: Optional[UploadFile] = File(None),
    archive: Optional[UploadFile] = File(None),
    output_filename: Optional[str] = Form("merged_document"),
    request: Request = None
):
    """API endpoint to merge files from a ZIP or RAR archive"""
    # Get the actual file from any of the possible parameter names
    actual_file = file or data or archive
    
    # Debug information
    debug_info = {
        "file_provided": file is not None,
        "data_provided": data is not None,
        "archive_provided": archive is not None,
        "output_filename": output_filename,
    }
    
    if not actual_file:
        # Try to get the file from form data directly
        form = await request.form()
        for key, value in form.items():
            if isinstance(value, UploadFile):
                actual_file = value
                debug_info["found_in_form"] = key
                break
    
    if not actual_file:
        raise HTTPException(status_code=400, detail=f"No file provided. Debug info: {debug_info}")
    
    # Asumimos que el archivo es un archivo comprimido válido y lo procesamos directamente
    # sin verificar la extensión, ya que n8n puede estar enviando archivos sin extensión
    
    # Create a temporary directory to store and extract files
    with tempfile.TemporaryDirectory() as temp_dir:
        # Save the uploaded archive with a nombre genérico para evitar problemas con extensiones
        temp_file_path = os.path.join(temp_dir, "uploaded_archive")
        with open(temp_file_path, "wb") as buffer:
            shutil.copyfileobj(actual_file.file, buffer)
        
        # Create extraction directory
        extract_dir = os.path.join(temp_dir, "extracted")
        os.makedirs(extract_dir, exist_ok=True)
        
        # Extract files from the archive
        try:
            extracted_files = extract_compressed_file(temp_file_path, extract_dir)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error extracting archive: {str(e)}")
        
        # Filter for PDF and DOCX files only
        valid_extensions = ['.pdf', '.docx', '.doc']
        filtered_files = filter_files_by_extension(extracted_files, valid_extensions)
        
        if not filtered_files:
            raise HTTPException(status_code=400, detail="No PDF or DOCX/DOC files found in the archive")
        
        # Check if all files are of the same type
        file_extensions = [os.path.splitext(path)[1].lower() for path in filtered_files]
        unique_extensions = set(file_extensions)
        
        if len(unique_extensions) > 1:
            raise HTTPException(status_code=400, 
                              detail="Files in the archive must be of the same type (either all PDF or all DOCX/DOC)")
        
        # Sort extracted files by part number in filename
        sorted_paths = sorted(filtered_files, key=lambda path: extract_part_number(os.path.basename(path)))
        
        # Determine file type and set output path
        file_ext = os.path.splitext(sorted_paths[0])[1].lower()
        output_path = f"uploads/{output_filename}{file_ext}"
        
        if file_ext == '.pdf':
            # For PDF files, merge them directly
            merge_pdf_files(sorted_paths, output_path)
        else:  # .docx or .doc
            # For DOCX files, merge them
            merge_docx_files_custom(sorted_paths, output_path)
        
        # Set the appropriate media type based on file extension
        if output_path.endswith(".pdf"):
            media_type = "application/pdf"
        else:  # .docx or .doc
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        return FileResponse(
            path=output_path,
            filename=os.path.basename(output_path),
            media_type=media_type
        )

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    uvicorn.run(app, host="0.0.0.0", port=port)
